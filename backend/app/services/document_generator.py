from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_resume_docx(resume_data: dict, output_path: str) -> str:
    doc = Document()

    # Set base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Name (header)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)

    # Contact info
    contact = resume_data.get("contact", {})
    contact_line = " | ".join(filter(None, [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("linkedin", ""),
        contact.get("github", "")
    ]))
    contact_para = doc.add_paragraph(contact_line)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_heading(text):
        heading = doc.add_paragraph()
        run = heading.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        heading.space_before = Pt(10)
        heading.space_after = Pt(4)

    # Summary
    if resume_data.get("summary"):
        add_section_heading("Professional Summary")
        doc.add_paragraph(resume_data["summary"])

    # Skills
    if resume_data.get("skills"):
        add_section_heading("Technical Skills")
        doc.add_paragraph(", ".join(resume_data["skills"]))

    # Projects
    if resume_data.get("projects"):
        add_section_heading("Projects")
        for project in resume_data["projects"]:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(project.get("name", ""))
            proj_run.bold = True

            if project.get("technologies"):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run("Technologies: " + ", ".join(project["technologies"]))
                tech_run.italic = True
                tech_run.font.size = Pt(9.5)

            if project.get("description"):
                doc.add_paragraph(project["description"])

    # Experience
    if resume_data.get("experience"):
        add_section_heading("Experience")
        for exp in resume_data["experience"]:
            exp_para = doc.add_paragraph()
            title_run = exp_para.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True

            if exp.get("duration"):
                duration_para = doc.add_paragraph()
                duration_run = duration_para.add_run(exp["duration"])
                duration_run.italic = True
                duration_run.font.size = Pt(9.5)

            if exp.get("description"):
                doc.add_paragraph(exp["description"])

    # Education
    if resume_data.get("education"):
        add_section_heading("Education")
        for edu in resume_data["education"]:
            edu_line = f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('year', '')})"
            doc.add_paragraph(edu_line)

    # Certifications
    if resume_data.get("certifications"):
        add_section_heading("Certifications")
        for cert in resume_data["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    # Set page margins (ATS-friendly: simple, no columns)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    doc.save(output_path)
    return output_path