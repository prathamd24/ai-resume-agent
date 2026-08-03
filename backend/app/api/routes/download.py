from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import docx
import tempfile
import os

router = APIRouter()

# This tells FastAPI exactly what JSON structure to expect from the frontend
class ResumeDownloadRequest(BaseModel):
    name: str
    email: str = ""
    phone: str = ""
    optimized_resume: dict

@router.post("/download")
async def generate_docx(data: ResumeDownloadRequest):
    """
    Takes the rewritten resume JSON from the frontend and converts it into a 
    downloadable Microsoft Word document.
    """
    try:
        # Create a new blank Word Document
        doc = docx.Document()
        
        # 1. Add Header (Name and Contact)
        doc.add_heading(data.name, 0)
        contact_info = f"{data.email} | {data.phone}"
        if contact_info.strip() != "|":
            doc.add_paragraph(contact_info)
            
        # 2. Add Professional Summary
        if "summary" in data.optimized_resume:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(data.optimized_resume["summary"])
            
        # 3. Add Experience
        if "experience" in data.optimized_resume:
            doc.add_heading("Experience", level=1)
            for exp in data.optimized_resume["experience"]:
                # Bold Title and Company
                p = doc.add_paragraph()
                title = exp.get('title', 'Role')
                company = exp.get('company', 'Company')
                p.add_run(f"{title} at {company}").bold = True
                
                # Add the AI-rewritten bullet points!
                for bullet in exp.get("optimized_bullet_points", []):
                    doc.add_paragraph(bullet, style='List Bullet')
                    
        # Save to a temporary file on the server
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        
        # Return the file as an actual download to the user's browser!
        return FileResponse(
            path=temp_file.name,
            filename=f"{data.name.replace(' ', '_')}_Optimized_Resume.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
